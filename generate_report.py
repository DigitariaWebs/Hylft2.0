from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

title = doc.add_heading('Daily Client Report', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_p.add_run('Date: 2026-05-17')
date_run.italic = True
date_run.font.size = Pt(11)

doc.add_heading('Client Conversations', level=1)

clients = [
    (
        'Maroin (Hylift)',
        'Meeting recorded. He gave me small enhancements; I implemented them, '
        'built the app, and sent him an APK.'
    ),
    (
        'Azzedin (La Menagere)',
        'Meeting recorded. Showed him the new app and the CRM. He gave a few '
        'small frontend improvements; I will start applying them the next '
        'working day and also begin implementing the backend.'
    ),
    (
        'Abdollah (Pops)',
        'We configured the VPS. I will ask Arslene to prepare a TestFlight build for him.'
    ),
    (
        'Patrick',
        'Scheduled a meeting for this Monday.'
    ),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
table.autofit = False

col_widths = (Cm(4.5), Cm(12.0))

hdr_cells = table.rows[0].cells
hdr_cells[0].text = ''
hdr_cells[1].text = ''
headers = ('Client', 'Conversation / Status')
for idx, (cell, label) in enumerate(zip(hdr_cells, headers)):
    cell.width = col_widths[idx]
    shade_cell(cell, '1F4E78')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(11)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

for name, convo in clients:
    row = table.add_row().cells
    row[0].width = col_widths[0]
    row[1].width = col_widths[1]

    name_p = row[0].paragraphs[0]
    name_run = name_p.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(11)
    row[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    convo_p = row[1].paragraphs[0]
    convo_run = convo_p.add_run(convo)
    convo_run.font.size = Pt(11)
    row[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP

doc.add_paragraph()

doc.add_heading('Report of the Day', level=1)

bullets = [
    'Analysed the document with Arslene across several separate sessions, '
    'fitting them between client meetings.',
    'Most of the day was spent in meetings with clients.',
    'Implemented the frontend improvements Maroin requested, built the app, '
    'and sent him an APK.',
]

for item in bullets:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(11)

output_path = r'c:\Users\Mohamed\Desktop\webProjects\Hylft2.0\Daily_Client_Report_2026-05-17.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
